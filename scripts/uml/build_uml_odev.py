#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""10 senaryo için UML tasarım çözümleri (referans ödev formatında) DOCX üretici.
Tümü DİKEY (portrait) A4, tüm kenarlardan 0,5 cm boşluk. Sınıf diyagramı dahildir;
sıra (sequence) diyagramı bilinçli olarak kaldırılmıştır.
Diyagramlar PlantUML ile scripts/uml/png/ altında üretilmiştir."""
import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "notes"))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from PIL import Image

PNG = os.path.join(os.path.dirname(__file__), "png")
import docx_helper as H


def _portrait(sec):
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(0.5)
    sec.top_margin = sec.bottom_margin = Cm(0.5)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    H._set_run(p.add_run(text), size=9, italic=True, color=H.GREY)


def add_diagram(doc, name, heading, cap):
    """Diyagramı dikey (portrait) sayfa akışına yerleştirir (tüm kenarlardan 0,5 cm).
    Bölüm/oryantasyon değişimi YOK — belge tamamen dikeydir."""
    H.h2(doc, heading)
    path = os.path.join(PNG, name)
    w_px, h_px = Image.open(path).size
    max_w, max_h = 20.0, 24.0  # cm (portrait içerik kutusu: 0,5 cm kenar + başlık/alt yazı payı)
    scale = min(max_w / w_px, max_h / h_px)
    doc.add_picture(path, width=Cm(w_px * scale), height=Cm(h_px * scale))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(doc, cap)


def numbered(doc, items):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(1)
        H._set_run(p.add_run(f"{i}. "), bold=True, color=H.ACCENT)
        H._set_run(p.add_run(it))


QUESTIONS = [
    {
        "no": 1,
        "baslik": "Personel Takvimi — Etkinlik Saati Değişimi",
        "patterns": "Observer + Mediator",
        "puan": 0,
        "analiz": [
            "Bir personelin takvimi ardışık etkinlikler serisidir; her etkinlik ad, tarih, salon no, başlangıç/bitiş saati ve diğer katılımcı personelin biricik numaralarını saklar.",
            "Bir etkinliğin saati değişince bağımlı taraflar — aynı gün SONRAKİ etkinliklerin katılımcıları ve toplantı salonu randevu sistemi — otomatik haberdar edilmeli; ancak etkinlik kime bildirim gittiğini bilmek zorunda olmamalı (yayıncı–abone ayrışması).",
            "Bildirimin birden çok hedefi var (çok katılımcı + randevu sistemi) ve hedefler zamanla değişebilir.",
            "Personellerin birbirini doğrudan tanıması n×(n-1) bağ üretir; salon randevu çakışmaları ise tek noktadan koordinasyon ister.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Observer", "Etkinlik «Subject»; KatilimciBildirimi ve SalonRandevuSistemi «Observer»",
             "Saat değişince bağımlılara otomatik bildirim; etkinlik somut alıcıya bağlanmaz, yeni bildirim kanalı eklemek etkinliği değiştirmez."),
            ("Mediator", "SalonRandevuSistemi",
             "Personeli birbirine doğrudan bağlamadan salon randevularını ve çakışmaları tek merkezden koordine eder."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("Doğrudan referans (etkinlik → her katılımcı)", "Sıkı bağ; yeni katılımcı tipi/kanal eklemek etkinliği değiştirir, izole test edilemez."),
            ("Polling (katılımcılar periyodik sorgular)", "Gecikme ve gereksiz yük; 'otomatik anında güncelleme' şartını karşılamaz."),
            ("Singleton tek küresel takvim", "Çok-personel/çok-istemci senaryosuna uymaz; test ve izolasyon zorlaşır."),
        ],
        "png": "soru1.png",
        "cap": "Şekil 1 · Personel Takvimi — Observer (bildirim) + Mediator (salon koordinasyonu).",
        "akis": [
            "Personel bir etkinliğin başlangıç/bitiş saatini değiştirir → Etkinlik.saatGuncelle().",
            "Etkinlik kendi durumunu günceller ve haberVer() çağırır.",
            "Takvim, ayniGunSonrakiler() ile aynı gün devamındaki etkinlikleri belirler.",
            "Her abone (katılımcı bildirimleri + salon randevu sistemi) guncelle() ile otomatik bilgilendirilir.",
            "SalonRandevuSistemi (Mediator) yeni saatte salon çakışması olup olmadığını denetler.",
        ],
        "kanaat": "Observer 'kim haberdar olacak' bilgisini etkinlikten ayırır; Mediator 'salon nasıl paylaşılacak' kuralını merkezîleştirir. İkisi birbirini tamamlar.",
    },
    {
        "no": 2,
        "baslik": "Ders Haftalık Planı — Teorik/Uygulamalı + Sınav Haftaları",
        "patterns": "Factory Method + Template Method",
        "analiz": [
            "Her ders biricik kod + ad/kredi/amaç/içerik özeti taşır ve istenen hafta sayısı kadar haftalık plan üretilebilmelidir.",
            "Plan üretim ALGORİTMASI sabittir: ara sınav sayısı = yuvarla(hafta/7), sınavlar eşit aralıkla yerleştirilir, kalan haftalar içerikle doldurulur. İskelet değişmez, türe-özgü adım değişir → Template Method.",
            "İçerik haftası türe göre farklı bir nesnedir (Teorik: konu başlığı + değerlendirme biçimi; Uygulamalı: uygulama özeti + beklenen çıktılar). Nesnenin yaratımı türe bırakılmalı → Factory Method.",
            "Sınav haftaları her iki türde ORTAKTIR (hafta no, sınav biçimi, süre, max puan).",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Template Method", "Ders.planUret()",
             "Sabit plan-üretim iskeleti (sınav sayısı + eşit aralık + doldurma); türe-özgü adımlar alt sınıfta override edilir."),
            ("Factory Method", "Ders.icerikHaftasiOlustur(); TeorikDers / UygulamaliDers",
             "Hangi HaftaOgesi (Teorik/Uygulama) nesnesinin yaratılacağını alt tür belirler; planUret() somut hafta sınıfına bağlanmaz."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("if/switch (ders türüne göre)", "Yeni tür eklemek planUret()'i değiştirir (OCP ihlali); tür kodu her yere dağılır."),
            ("Abstract Factory", "Tek ürün ailesi (hafta) var; tam aile gerektirmediği için aşırı tasarım — Factory Method yeterli."),
            ("Strategy (plan algoritması dışarıdan)", "Algoritma sabit; değişen yalnız hafta nesnesinin tipi → Factory Method daha doğru eşleşir."),
        ],
        "png": "soru2.png",
        "cap": "Şekil 2 · Ders Haftalık Planı — Template Method (algoritma) + Factory Method (hafta nesnesi).",
        "akis": [
            "Ders tanımlanır (kod/ad/.../haftaSayisi) ve tür seçilir (Teorik/Uygulamalı).",
            "planUret(hafta) çağrılır → araSinavSayisi = yuvarla(hafta/7).",
            "Sınav haftaları eşit aralıkla SinavHaftasi olarak yerleştirilir (her iki türde ortak).",
            "Kalan haftalar icerikHaftasiOlustur() ile türe göre TeorikHafta / UygulamaHaftasi olarak doldurulur.",
            "Tür-uygun HaftalikPlan döner.",
        ],
        "kanaat": "Değişmeyeni (algoritma) Template Method'da sabitle, değişeni (hafta nesnesi) Factory Method'a devret.",
    },
    {
        "no": 3,
        "baslik": "Oyun — Savunma/Saldırı Araçları (İlkel/Gelişkin Düzey)",
        "patterns": "Abstract Factory + Strategy",
        "analiz": [
            "Kahraman'ın bir saldırı + bir savunma aracı vardır; düzeyler araçları AİLE olarak belirler: Düzey-1 {kılıç, kalkan}, Düzey-2 {tabanca, çelik yelek}. Aile içi tutarlılık (ilkel saldırının yanında ilkel savunma) korunmalı → Abstract Factory.",
            "koruBeni() ve saldirSimdi() iletileri ÇOK BİÇİMLİ kullanılır → araç davranışları değiştirilebilir algoritmalardır → Strategy.",
            "Kahraman somut araç sınıflarına (Kılıç, Tabanca...) bağlanmamalı; düzey değişince donanım tutarlı şekilde yenilenmelidir.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Abstract Factory", "AracFabrikasi; IlkelAracFabrikasi / GeliskinAracFabrikasi",
             "Birbiriyle uyumlu saldırı+savunma çiftini tek elden üretir; düzey değişince tüm aile tutarlı değişir, uyumsuz çift oluşmaz."),
            ("Strategy", "SaldiriAraci / SavunmaAraci arayüzleri",
             "saldirSimdi() / koruBeni() çok biçimli; Kahraman somut araca bağlanmadan davranışı çalışma anında değiştirebilir."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("if/switch (düzeye göre araç seçimi)", "Yeni düzey eklemek Kahraman'ı değiştirir; uyumsuz çift (ilkel saldırı + gelişkin savunma) riski doğar."),
            ("Factory Method", "Tek ürün değil EŞLEŞMİŞ aile (saldırı+savunma) gerekiyor → Abstract Factory daha doğru."),
            ("Kalıtımla araç gömme (her düzey için ayrı Kahraman)", "Davranış sabitlenir; çalışma anında araç değişimi engellenir."),
        ],
        "png": "soru3.png",
        "cap": "Şekil 3 · Oyun Araçları — Abstract Factory (uyumlu aile) + Strategy (çok biçimli iletiler).",
        "akis": [
            "Düzeye göre uygun fabrika seçilir (IlkelAracFabrikasi / GeliskinAracFabrikasi).",
            "Kahraman.donanimYukle(fabrika) → saldiriYarat() + savunmaYarat() uyumlu çifti döndürür.",
            "Oyun sırasında saldirSimdi() aktif SaldiriAraci'ya, koruBeni() aktif SavunmaAraci'ya çok biçimli yönlenir.",
            "Düzey atlanınca fabrika değişir; tüm aile tutarlı biçimde yenilenir.",
        ],
        "kanaat": "Abstract Factory 'hangi aile' tutarlılığını, Strategy 'davranış nasıl çağrılır' esnekliğini verir.",
    },
    {
        "no": 4,
        "baslik": "Kütüphane — Kitap Ödünç Takibi",
        "patterns": "Alan Modeli (İlişki Sınıfı) — GoF örüntüsü gerekmez",
        "analiz": [
            "Bir kitap farklı zamanlarda farklı öğrenciler tarafından ödünç alınır; bir öğrenci aynı anda birden çok kitap alabilir → çok-çoğa ilişki.",
            "Hangi kitabın kim tarafından alındığı ve ne zaman iade edildiği izlenmeli → ilişkinin kendi öznitelikleri (alış/iade tarihi) vardır.",
            "Bu bir davranış/üretim/erişim sorunu DEĞİL, saf veri-ilişki modellemesidir; GoF örüntüsü zorlamak yanlış olur.",
        ],
        "secim_h": ("Modelleme Kararı", "Uygulanan Yer", "Gerekçe"),
        "secim": [
            ("İlişki Sınıfı (Association Class)", "Odunc",
             "Çok-çoğa ilişkiyi alış/iade tarihiyle birlikte taşır; her ödünç olayı ayrı kayıttır."),
            ("Çokluk", "Ogrenci 1—0..* Odunc 0..*—1 Kitap",
             "Zaman içinde tekrar eden ödünç kayıtlarını ve aynı anda çok kitabı destekler."),
        ],
        "red_t": "Zorlanabilecek ama Uygun Olmayan Örüntüler",
        "red_h": ("Örüntü", "Neden gereksiz?"),
        "red": [
            ("State (kitap: rafta/ödünçte)", "Tek bayrak yeterli; ayrı durum sınıfları aşırı. iadeTarihi = boş zaten 'ödünçte'yi ifade eder."),
            ("Observer (gecikmiş iade bildirimi)", "Senaryoda istenmiyor; eklenirse kapsam-dışı varsayım olur (zero-hallucination)."),
        ],
        "png": "soru4.png",
        "cap": "Şekil 4 · Kütüphane — Ogrenci–Kitap çok-çoğa ilişkisi, Odunc ilişki sınıfı.",
        "akis": [
            "Öğrenci kitabı ödünç alır → yeni Odunc {alisTarihi = bugün, iadeTarihi = boş}.",
            "Aynı öğrenci başka kitap alır → ikinci Odunc kaydı (aynı anda çok kitap).",
            "Kitap iade edilince ilgili Odunc.iadeTarihi doldurulur.",
            "Geçmiş tüm ödünç kayıtları korunur (bir kitap zamanla çok öğrenciye gitmiş olabilir).",
        ],
        "kanaat": "Doğru çözüm 'örüntü' değil, doğru sınıf + çokluk + ilişki sınıfıdır; gereksiz örüntü yalnızca karmaşa üretir.",
    },
    {
        "no": 5,
        "baslik": "Sinema — Film/Salon/Seans/Bilet/Koltuk",
        "patterns": "Alan Modeli — GoF örüntüsü gerekmez",
        "analiz": [
            "Farklı filmler farklı salonlarda belirli saatlerde gösterilir; bir film çok seansa sahiptir; her seans tek salonda gerçekleşir.",
            "Müşteriler bir seans için bilet alır; bir seansa çok bilet satılabilir; her bilet tek bir seans + tek bir koltuk için kesilir.",
            "Bir koltuk farklı seanslarda yeniden satılır; bu yüzden Koltuk—Bilet bağı 1—0..* olmalı, 1—1 değil. Aynı (Seans, Koltuk) çifti en çok bir bilete bağlanır.",
            "Saf alan modellemesi; davranışsal bir örüntüye ihtiyaç yoktur — çözüm doğru sınıf ve çokluklardadır.",
        ],
        "secim_h": ("Modelleme Kararı", "Uygulanan Yer", "Gerekçe"),
        "secim": [
            ("Çokluk", "Film 1—0..* Seans; Salon 1—0..* Seans",
             "Bir film/salon birçok seans barındırır; her seans tek film ve tek salona bağlıdır."),
            ("Kompozisyon", "Salon ◆—1..* Koltuk",
             "Koltuk salona aittir; salon yoksa koltuk da yoktur (güçlü yaşam-döngüsü bağı)."),
            ("Bilet–Koltuk", "Koltuk 1—0..* Bilet; Seans 1—0..* Bilet",
             "Bir koltuk farklı seanslarda çok bilete girer; her bilet tek koltuk + tek seans içindir. Kısıt: aynı (Seans, Koltuk) en çok bir bilet."),
        ],
        "red_t": "Zorlanabilecek ama Uygun Olmayan Örüntüler",
        "red_h": ("Örüntü", "Neden gereksiz?"),
        "red": [
            ("Singleton (BiletSistemi)", "Çoklu salon/seans için gereksiz küresel durum; test ve izolasyon zorlaşır."),
            ("Strategy (fiyatlandırma)", "Senaryoda fiyat kuralı tanımlı değil; eklemek kapsam-dışı varsayım olur."),
        ],
        "png": "soru5.png",
        "cap": "Şekil 5 · Sinema Bilet Sistemi — çokluklar ve Salon–Koltuk kompozisyonu.",
        "akis": [
            "Film + salon + tarih/saat ile bir Seans oluşturulur.",
            "Salon koltuk düzeniyle (1..* Koltuk) tanımlıdır.",
            "Müşteri bir seans seçer ve boş bir koltuk için Bilet alır.",
            "Aynı seansa farklı koltuklar için çok sayıda Bilet satılabilir.",
        ],
        "kanaat": "Net çokluk ve kompozisyon ilişkileri doğru kurulduğunda model eksiksizdir; örüntü eklemek değer katmaz.",
    },
    {
        "no": 6,
        "baslik": "Kafe — İçecek Eklentileri ve İndirimler",
        "patterns": "Decorator (iki katman)",
        "analiz": [
            "Temel içecek (Espresso/Americano/Filtre Kahve) üzerine süt, ekstra shot, vanilya, karamel İSTENİLDİĞİ KADAR eklenir; her eklemeden sonra toplam fiyat ve ürün açıklaması üretilmeli → çalışma anında yinelemeli ekleme → Decorator.",
            "Ödemede müşterinin sahip olduğu indirimlerin (öğrenci + sadakat + kampanya) HEPSİ birlikte uygulanabilir → indirimlerin üst üste sarmalanması → ikinci Decorator katmanı.",
            "Her eklenti/indirim aynı arayüzü uygular; bu sayede kombinasyon patlaması olmadan sınırsız bileşim kurulur.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Decorator (eklenti)", "Icecek; EklentiDecorator → Sut/EkstraShot/Vanilya/Karamel",
             "Her eklenti aynı Icecek arayüzünü sarmalar; fiyat()/aciklama() özyinelemeli birikir → sınırsız ekleme, her adımda güncel toplam."),
            ("Decorator (indirim)", "IndirimDecorator → OgrenciIndirimi/SadakatIndirimi/KampanyaIndirimi",
             "Sahip olunan indirimlerin hepsi üst üste sarmalanarak fiyat()'a uygulanır."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("Alt sınıf patlaması (EspressoSutluVanilyali...)", "Her kombinasyon için ayrı sınıf → üstel artış; bakımı imkânsız."),
            ("Strategy (tek eklenti seçimi)", "'İstenildiği kadar ve birlikte' çoklu birikimi tek strateji ifade edemez."),
            ("if/switch ile fiyat toplama", "Yeni eklenti/indirim eklemek merkezî kodu değiştirir (OCP ihlali)."),
        ],
        "png": "soru6.png",
        "cap": "Şekil 6 · Kafe — içecek eklentileri ve indirimler için iki katmanlı Decorator.",
        "akis": [
            "Temel içecek seçilir (ör. Espresso).",
            "Eklenti sarmalanır: new Sut(new EkstraShot(espresso)) → her adımda aciklama()/fiyat() güncel toplamı verir.",
            "Ödemede indirimler sarmalanır: new KampanyaIndirimi(new OgrenciIndirimi(icecek)).",
            "Son fiyat() çağrısı tüm eklenti + indirim zincirini özyinelemeli hesaplar.",
        ],
        "kanaat": "Decorator 'sınırsız ve birlikte' birikimi sınıf patlaması olmadan modeller; aynı arayüz hem eklemede hem indirimde yeniden kullanılır.",
    },
    {
        "no": 7,
        "baslik": "Savunma Sanayi — Belge Hazırlama/Onay/Yayın Yönetimi",
        "patterns": "State + Chain of Responsibility + Strategy",
        "analiz": [
            "Her belge hazırlama → onaylama → yayınlama işlemlerine sahiptir; belge 'hazırlık aşamasında / X işlemde' biçiminde bir yaşam döngüsü izler ve her durumda yalnız uygun işlem geçerlidir → State.",
            "Her işlemi ilkesel olarak farklı kullanıcı/birim yapar; ONAY makamı belge türü ve gizlilik derecesine göre değişir (Orta → Birim Yöneticisi, 2 imza; Yüksek → Koordinatörlük, tüm birim yöneticilerinin imzası) → istek uygun makamı bulana dek iletilir → Chain of Responsibility.",
            "Milli Gizli belgelerin HAZIRLAMA işlemi dereceye göre farklıdır ve derece DEĞİŞEBİLİR → çalışma anında değiştirilebilir davranış → Strategy.",
            "Belge türleri Milli Gizli / Genel; Genel belgeler Halk (ilan seri-sıra no) ve Kamu Kurumları (Devlet Protokolü no + damga) olarak ayrışır → yayınlama farkı alt tür polimorfizmiyle çözülür.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("State", "BelgeDurumu; HazirlikAsamasi / OnayAsamasi / YayinAsamasi",
             "'Hazırlık → onay → yayın' akışını yönetir; geçersiz işlem (ör. hazırlanmamış belgeyi yayınlamak) durum tarafından engellenir."),
            ("Chain of Responsibility", "OnayMakami; BirimYoneticisi → Koordinatorluk",
             "Onay isteği zincirde ilerler; Orta gizli Birim Yöneticisinde (2 imza), Yüksek gizli Koordinatörlükte (tüm birim yön. imzası) karşılanır."),
            ("Strategy", "HazirlamaStratejisi; OrtaGizliHazirlama / YuksekGizliHazirlama",
             "Dereceye göre farklı hazırlama; derece değişince strateji değişir, Belge sınıfı değişmez."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("if/switch (her işlemde tür+derece)", "Onay/hazırlama/yayın kuralları her yere dağılır; yeni makam/derece eklemek çok noktayı kırar (OCP)."),
            ("Tek dev Belge sınıfı (bayraklarla)", "Durum + tür + derece kombinasyonu kombinasyon patlamasına yol açar; okunamaz hale gelir."),
            ("Observer (onay bildirimi)", "Asıl sorun yetki YÖNLENDİRME (kim onaylar), bildirim değil; CoR daha doğru eşleşir."),
        ],
        "png": "soru7.png",
        "cap": "Şekil 7 · Belge Yönetimi — State (akış) + CoR (onay zinciri) + Strategy (dereceye göre hazırlama). Belge türleri kalıtımla, yayınlama farkı alt tür polimorfizmiyle.",
        "akis": [
            "Belge oluşturulur; durum = HazirlikAsamasi. Tür (Milli Gizli/Genel) ve gerekirse derece belirlenir.",
            "hazirla(): Milli Gizli ise dereceye uygun HazirlamaStratejisi (Orta/Yüksek) çalışır; durum OnayAsamasi'na geçer.",
            "onayla(): istek OnayMakami zincirine girer; uygun makam (Birim Yöneticisi / Koordinatörlük) gerekli imzaları toplar.",
            "yayinla(): tür alt sınıfına göre Halk (ilan seri-sıra no) ya da Kamu (Devlet Protokolü no + damga); durum YayinAsamasi'na geçer.",
        ],
        "kanaat": "State zamanı (akış aşaması), CoR yetkiyi (kim onaylar), Strategy türe-özgü işlemi ayırır. Üçü birbirinin yükünü devralmaz; biri çıkarılırsa diğer ikisi şişer.",
    },
    {
        "no": 8,
        "baslik": "Akıllı Cihazlar — Güç Eşiğine Göre Uygulama Yönetimi",
        "patterns": "Observer + State + Memento",
        "analiz": [
            "Cihazlarda (Akıllı Saat / Telefon / Tablet) uygulamalar kapatılmadıkça açıktır; güç seviyesi değişince ilgili uygulamalar haberdar edilmeli, ancak uygulama cihaza doğrudan bağlı olmamalı → Observer.",
            "Eşik davranışı: %20 altı konum servisi, %15 altı bluetooth, %10 altı hareket izleme uygulamaları cihazla etkileşimi keser (güç-tasarrufu); şarjda iken seviye geri çıkınca normale döner → uygulamanın iki davranış durumu → State.",
            "'En son o seviyedeyken kesilmiş' uygulamanın 'normal-beklenen' haline geri dönmesi → kesilmeden önceki durumun saklanıp geri yüklenmesi → Memento.",
            "Geri dönüş AŞAMALIDIR: şarjda seviye %10 → %15 → %20 üstüne çıktıkça, o seviyede kesilmiş uygulamalar sırayla normale döner.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Observer", "Cihaz «Subject»; Uygulama «Observer» (GucGozlemcisi)",
             "gucDegisti() → haberVer(); seviye değişince eşiği aşan uygulamalara bildirim; cihaz somut uygulamaya bağlanmaz."),
            ("State", "UygulamaDurumu; NormalDurum / GucTasarrufuDurumu",
             "Eşik altına inince GucTasarrufu, şarjda seviye eşik üstüne çıkınca Normal; davranış if/switch olmadan değişir."),
            ("Memento", "UygulamaAnisi",
             "Kesilmeden önce anaKaydet(); seviye gelince aniGeriYukle() ile son duruma dönüş; iç durum kapsüllenmiş kalır."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("Polling (uygulama gücü sürekli sorgular)", "Sürekli yoklama enerji harcar — düşük güç senaryosuyla çelişir; Observer push daha doğru."),
            ("if/switch (güç bandına göre)", "Üç eşik + iki yön (düşüş/şarj) çok dallı ve kırılgan; State temiz çözer."),
            ("Durumu uygulama dışında (global) tutmak", "Kapsülleme bozulur; Memento iç durumu açığa çıkarmadan saklar/geri yükler."),
        ],
        "png": "soru8.png",
        "cap": "Şekil 8 · Akıllı Cihaz Güç Yönetimi — Observer (yayın) + State (davranış) + Memento (geri yükleme).",
        "akis": [
            "Uygulamalar Cihaz'a abone olur (Observer); her uygulamanın eşiği vardır (konum %20 / bluetooth %15 / hareket %10).",
            "Güç düşer, eşik aşılır → Cihaz haberVer() → ilgili uygulama anaKaydet() ile durumunu saklar (Memento) ve GucTasarrufuDurumu'na geçip etkileşimi keser (State).",
            "Cihaz enerjiye bağlanır; seviye %10 → %15 → %20 üstüne aşamalı çıkar.",
            "Her eşik geçişinde o seviyede kesilmiş uygulamalara bildirim → aniGeriYukle() + NormalDurum → normal-beklenen çalışmaya döner.",
        ],
        "kanaat": "Observer haberleşmeyi, State davranış değişimini, Memento 'geri dönüş noktasını' ayırır; üçü birlikte 'aşamalı kes / geri-yükle' döngüsünü eksiksiz kurar.",
    },
    {
        "no": 9,
        "baslik": "Akıllı Ev — Bağlantılı Çalışma Modları ve Sensör Bildirimleri",
        "patterns": "State + Mediator + Observer",
        "analiz": [
            "Aydınlatma ve ısıtma sistemleri ile kapı ve hareket sensörleri verilerini akıllı ev izleme motoruna aktarır ve motordan gelen talimatları yerine getirir.",
            "Aydınlatma ve ısıtmanın güvenli ve tasarruflu iki çalışma şekli vardır ve her şekilde davranış farklıdır (tasarruflu aydınlatma: ortam ışığına destekle asgari aydınlatma; güvenli: en yüksek kapasite. Tasarruflu ısıtma: ortamı 18°C'ye çıkarır; güvenli: her durumda 40°C panel ısısı) → nesnenin içinde bulunduğu duruma göre davranışı değişir → State.",
            "İki çalışma şekli BİRBİRİYLE BAĞLANTILIDIR: aydınlatma tasarrufluya geçerse ısıtma da tasarrufluya geçmeli (ya da tam tersi). Sistemler birbirini doğrudan tanımadan tek merkezden eşitlenmeli → Mediator (izleme motoru).",
            "Ev sahibi mobil/web uygulamasından tüm sistem ve sensör durumlarını görür ve izler; durum değişince haberdar edilmeli → Observer.",
            "Kapı ve hareket sensörleri için, ETKİNLEŞTİRİLMİŞSE, tetiklenmede bildirim doğrudan kullanıcının uygulamasına SMS olarak yönlendirilmelidir → Observer bildiriminin bir kanalı.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("State", "CalismaModu; TasarrufluAydinlatma/GuvenliAydinlatma, TasarrufluIsitma/GuvenliIsitma",
             "Her sistemin güvenli/tasarruflu davranışı (asgari↔azami aydınlatma; 18°C↔40°C) içinde bulunduğu moda göre değişir; if/switch'e gerek kalmadan."),
            ("Mediator", "AkilliEvMotoru",
             "İki çalışma şeklini bağlantılı tutar (biri tasarruflu→diğeri tasarruflu); sistemleri birbirine doğrudan bağlamadan koordine eder, sensör verisini alıp talimat gönderir."),
            ("Observer", "AkilliEvMotoru «Subject»; KullaniciUygulamasi «Observer»",
             "Kullanıcı tüm sistem/sensör durumlarını uygulamadan izler; kapı/hareket sensörü etkinse tetiklenmede doğrudan SMS bildirimi gönderilir."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("if/switch (moda göre her sistemde)", "İki mod × iki sistem dallanması her yere yayılır; yeni mod/sistem eklemek çok noktayı kırar (OCP). State temiz çözer."),
            ("Aydınlatma ↔ ısıtma doğrudan bağ", "Sistemler birbirini doğrudan tanırsa sıkı bağ oluşur; yeni sistem eklemek mevcutları değiştirir. Mediator bağlantıyı merkezîleştirir."),
            ("Polling (uygulama durumları periyodik sorgular)", "Gecikme + gereksiz yük; 'anında izleme/bildirim' şartını karşılamaz. Observer push daha doğru."),
            ("Strategy (mod = değiştirilebilir algoritma)", "Strategy'ye yakın; ancak mod, sistemin İÇİNDE BULUNDUĞU ve geçiş yaptığı bir durumdur (tasarruflu↔güvenli) → State daha doğru eşleşir."),
        ],
        "png": "soru9.png",
        "cap": "Şekil 9 · Akıllı Ev — State (çalışma modu) + Mediator (bağlantılı mod) + Observer (izleme/SMS).",
        "akis": [
            "Ev sahibi uygulamadan aydınlatmayı tasarruflu moda alır → AkilliEvMotoru.modDegistir(aydinlatma, TASARRUFLU).",
            "Motor aydınlatmanın modunu TasarrufluAydinlatma yapar; uygula() ortam ışığına destekle asgari aydınlatmayı sağlar (State).",
            "Bağlantı gereği motor ısıtmayı da TASARRUFLU yapar; TasarrufluIsitma.uygula() ortam sıcaklığını 18°C'ye çıkarır.",
            "Motor durum değişimini abonelere haber verir; kullanıcı uygulaması güncel durumu gösterir (Observer).",
            "Kapı/hareket sensörü tetiklenir → veriAktar(); motor durumu uygulamaya iletir. Sensörün bildirimi etkinse uygulamaya SMS gönderilir.",
        ],
        "kanaat": "State 'hangi modda nasıl davranılır'ı, Mediator 'iki modun bağlantılı kalması'nı, Observer 'kullanıcının izlemesi ve SMS bildirimi'ni ayırır; üçü birbirinin yükünü devralmaz.",
    },
    {
        "no": 10,
        "baslik": "SimpleCar E-Satış — Ödeme, Kontroller ve Sipariş İşleme",
        "patterns": "Strategy + Chain of Responsibility + Template Method",
        "analiz": [
            "SimpleCar tek model araba üretir ve yalnızca e-satış yapar; siparişler uygulamadan gelir.",
            "Ödeme yöntemi banka havalesi, PayPal ya da soğuk cüzdan olabilir ve bu yöntemler SIK genişler/daralır → çalışma anında değiştirilebilen, kolay eklenip çıkarılabilen davranış → Strategy.",
            "Her sipariş sahtecilik, limit ve bakiye kontrollerine tâbidir; kontrollerin HERHANGİ BİRİ olumsuzsa sipariş iptal edilir ve bu kontroller zamanla ARTIRILABİLİR → istek, her biri reddedebilen denetçiler zincirinden geçer → Chain of Responsibility.",
            "Kontroller başarılıysa sipariş işlenir: fatura düzenleme → fatura gönderme → alacaklar listesine kaydetme adımları SIRASIYLA tamamlanır → sabit sıralı iskelet → Template Method.",
        ],
        "secim_h": ("Örüntü", "Uygulanan Yer", "Çözdüğü Olgu"),
        "secim": [
            ("Strategy", "OdemeYontemi; BankaHavalesi/PayPal/SogukCuzdan",
             "Ödeme yöntemini çalışma anında seçer; sık genişleyen/daralan yöntemler Siparis sınıfını değiştirmeden eklenir/çıkarılır (OCP)."),
            ("Chain of Responsibility", "KontrolHalkasi; SahtecilikKontrolu → LimitKontrolu → BakiyeKontrolu",
             "Sipariş kontrolleri zincirde sırayla işler; biri olumsuz olursa sipariş iptal edilir; yeni kontrol zincire eklenerek artırılır."),
            ("Template Method", "SiparisIsleyici.siparisiIsle()",
             "Kontroller başarılıysa işleme iskeleti sabittir: fatura düzenle → fatura gönder → alacaklara kaydet; adımların sırası garanti edilir."),
        ],
        "red_t": "Reddedilen Alternatifler",
        "red_h": ("Alternatif", "Neden tercih edilmedi?"),
        "red": [
            ("if/switch (ödeme türüne göre)", "Sık genişleyen/daralan yöntemler merkezî kodu sürekli kırar (OCP ihlali); Strategy her yöntemi izole eder."),
            ("Tek dev kontrol metodu (iç içe if)", "Yeni kontrol eklemek metodu değiştirir; her kontrolün ayrı reddetme yetkisi/sırası kaybolur. CoR her halkayı bağımsız ekler/çıkarır."),
            ("Observer (kontrol başarısızlığı bildirimi)", "Asıl sorun yayın/bildirim değil, sıralı ve reddedilebilir DENETİM akışıdır; CoR daha doğru eşleşir."),
            ("Adımları satır içi gömmek (iskeletsiz)", "Fatura düzenle/gönder/kaydet sırası garanti edilmez ve yeniden kullanılamaz; Template Method sıralı iskeleti sabitler."),
        ],
        "png": "soru10.png",
        "cap": "Şekil 10 · SimpleCar E-Satış — Strategy (ödeme) + CoR (kontroller) + Template Method (işleme).",
        "akis": [
            "Müşteri e-satış uygulamasından sipariş verir (miktar + siparişi veren bilgisi + ödeme yöntemi).",
            "Sipariş kontrol zincirine girer: sahtecilik → limit → bakiye (Chain of Responsibility).",
            "Kontrollerden biri olumsuzsa sipariş iptal edilir (durum = IPTAL).",
            "Hepsi başarılıysa seçili OdemeYontemi.ode(tutar) ile ödeme alınır (Strategy).",
            "SiparisIsleyici.siparisiIsle() sabit sırayla faturayı düzenler, gönderir ve alacaklar listesine kaydeder (Template Method); durum = ISLENDI.",
        ],
        "kanaat": "Strategy 'nasıl ödenir'i, CoR 'sipariş geçerli mi / kim reddeder'i, Template Method 'geçerliyse hangi adımlar hangi sırayla'yı ayırır; biri çıkarılırsa diğerleri şişer.",
    },
]


def build():
    doc = H.new_doc()
    _portrait(doc.sections[0])
    H.title_block(
        doc,
        "Nesneye Yönelik Yazılım Geliştirme",
        "Tasarım Senaryoları · UML 2.x Sınıf Diyagramları",
        "10 Senaryo · Senaryo Çözümlemesi · Örüntü Seçimi · Sınıf Diyagramı · Çalışma Akışı",
    )
    H.para(
        doc,
        "Bu belge, on tasarım senaryosu için UML 2.x çözümlerini hocanın istediği biçime birebir uyarak sunar. "
        "Her senaryo için sırasıyla: (i) senaryo çözümlemesi, (ii) örüntü seçimi ve reddedilen alternatiflerin "
        "gerekçeleri, (iii) UML 2.x sınıf diyagramı, (iv) tipik çalışma akışı verilir. Diyagramlar elle tasarlanmış "
        "olup düşük bağ (coupling) ve yüksek tutarlılık (cohesion) ile okunaklılık gözetir. Tüm tasarım kararları "
        "senaryo metnine birebir sadıktır; veriye dayanmayan varsayım eklenmemiştir.",
    )
    H.para(
        doc,
        "Not: Bir senaryo birden çok örüntü gerektirebilir. 4 ve 5. sorular saf alan modellemesidir; bu sorularda "
        "GoF örüntüsü zorlamak yanlış olacağından bilinçli olarak örüntüsüz, doğru çokluk/ilişki temelli çözülmüştür.",
        italic=True, color=H.GREY,
    )

    H.h2(doc, "İçindekiler", color=H.ACCENT)
    H.table(
        doc,
        ("Soru", "Senaryo", "Örüntü / Yaklaşım"),
        [(f"Soru {q['no']}", q["baslik"], q["patterns"]) for q in QUESTIONS],
        widths=[1.6, 9.5, 6.4],
    )

    for qi, q in enumerate(QUESTIONS):
        doc.add_page_break()
        no = q["no"]
        H.h1(doc, q["baslik"], num=f"Soru {no}")
        H.para(doc, f"Kullanılan örüntü: {q['patterns']}", italic=True, color=H.ACCENT2, size=11.5)

        # (i) Senaryo çözümlemesi
        H.h2(doc, f"{no}.1  Senaryo Çözümlemesi")
        numbered(doc, q["analiz"])

        # (ii) Örüntü seçimi + reddedilen alternatif gerekçeleri
        H.h2(doc, f"{no}.2  Örüntü / Modelleme Seçimi")
        H.table(doc, q["secim_h"], q["secim"], widths=[3.4, 5.6, 8.8])
        H.h2(doc, f"{no}.3  {q['red_t']}")
        H.table(doc, q["red_h"], q["red"], widths=[6.0, 11.8])

        # (iii) UML 2.x sınıf diyagramı — dikey akış
        add_diagram(doc, q["png"], f"{no}.4  Sınıf Diyagramı (UML 2.x)", q["cap"])

        # (iv) Tipik çalışma akışı
        H.h2(doc, f"{no}.5  Tipik Çalışma Akışı")
        numbered(doc, q["akis"])
        H.callout(doc, "Kanaat:", q["kanaat"])

    out = os.path.join(os.path.dirname(__file__), "..", "..",
                       "attached_assets", "OOP_8_Senaryo_UML_Tasarim_Cozumleri.docx")
    out = os.path.abspath(out)
    doc.save(out)
    print("KAYDEDİLDİ:", out)
    return out


if __name__ == "__main__":
    build()
