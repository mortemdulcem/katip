#!/usr/bin/env python3
"""Black Hat Go Türkçe çeviri chunk'larını tek DOCX'e birleştirir.
Düz metin: Times New Roman 10pt. Kod: Consolas. Başlıklar korunur. A4 dar kenar, 1.15 satır."""
import os, re, glob
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CACHE = "scripts/blackhat/cache"
OUT = "client/public/BlackHatGo_Turkce.docx"

doc = Document()

# --- sayfa düzeni: A4, dar kenar (~1.27cm), 1.15 satır ---
sec = doc.sections[0]
sec.page_height = Inches(11.69)
sec.page_width = Inches(8.27)
for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
    setattr(sec, m, Inches(0.5))

# stil tabanı
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
def clean(s):
    return _CTRL.sub("", s)

def set_run_font(run, name, size=10, bold=False, italic=False, color=None):
    if run.text:
        run.text = clean(run.text)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(a), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def shade(cell_or_par, hexcolor):
    el = cell_or_par
    pPr = el.get_or_add_pPr() if hasattr(el, "get_or_add_pPr") else None

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")

def add_inline(p, text, base="Times New Roman", size=10):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            set_run_font(p.add_run(tok[2:-2]), base, size, bold=True)
        elif tok.startswith("`") and tok.endswith("`"):
            set_run_font(p.add_run(tok[1:-1]), "Consolas", size)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            set_run_font(p.add_run(tok[1:-1]), base, size, italic=True)
        else:
            set_run_font(p.add_run(tok), base, size)

def para_spacing(p):
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)

def add_heading(text, level):
    p = doc.add_paragraph()
    para_spacing(p)
    sizes = {1: 18, 2: 14, 3: 12, 4: 11}
    add_inline(p, text, size=sizes.get(level, 11))
    for r in p.runs:
        r.bold = True
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)

def add_code_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1)
    # gri arka plan
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    txt = "\n".join(lines)
    run = p.add_run(txt)
    set_run_font(run, "Consolas", 9)

def add_table(rows):
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            cp = cells[ci].paragraphs[0]
            para_spacing(cp)
            add_inline(cp, val.strip(), size=9)
            if ri == 0:
                for r in cp.runs:
                    r.bold = True

def is_table_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}.*\|", line)) and set(line.replace("|", "").replace(":", "").strip()) <= {"-", " "}

files = sorted(glob.glob(os.path.join(CACHE, "chunk_*.md")))
print(f"{len(files)} chunk birleştiriliyor...")

# başlık sayfası
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(tp.add_run("BLACK HAT GO"), "Times New Roman", 26, bold=True)
tp2 = doc.add_paragraph(); tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(tp2.add_run("Hacker’lar ve Sızma Testi Uzmanları için Go Programlama"), "Times New Roman", 14, italic=True)
tp3 = doc.add_paragraph(); tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(tp3.add_run("Tom Steele, Chris Patten, Dan Kottmann — No Starch Press, 2020"), "Times New Roman", 11)
tp4 = doc.add_paragraph(); tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_run_font(tp4.add_run("Türkçe çeviri"), "Times New Roman", 11)

for f in files:
    text = clean(open(f, encoding="utf-8").read())
    lines = text.split("\n")
    i = 0
    in_code = False
    code_lines = []
    tbl_buf = []
    def flush_tbl():
        global tbl_buf
        if tbl_buf:
            rows = []
            for ln in tbl_buf:
                if is_table_sep(ln):
                    continue
                cells = [c.strip() for c in ln.strip().strip("|").split("|")]
                rows.append(cells)
            add_table(rows)
            tbl_buf = []
    while i < len(lines):
        ln = lines[i]
        fence = re.match(r"^\s*```", ln)
        if fence:
            if not in_code:
                flush_tbl()
                in_code = True; code_lines = []
            else:
                add_code_block(code_lines); in_code = False; code_lines = []
            i += 1; continue
        if in_code:
            code_lines.append(ln); i += 1; continue
        # tablo satırı
        if ln.strip().startswith("|") and ln.count("|") >= 2:
            tbl_buf.append(ln); i += 1; continue
        else:
            flush_tbl()
        s = ln.strip()
        if not s:
            i += 1; continue
        h = re.match(r"^(#{1,6})\s+(.*)$", s)
        if h:
            add_heading(h.group(2), len(h.group(1))); i += 1; continue
        if re.match(r"^---+$", s) or re.match(r"^\*\*\*+$", s):
            i += 1; continue
        # liste
        lm = re.match(r"^[-*+]\s+(.*)$", s)
        nm = re.match(r"^(\d+)\.\s+(.*)$", s)
        if lm:
            p = doc.add_paragraph(style="List Bullet"); para_spacing(p); add_inline(p, lm.group(1)); i += 1; continue
        if nm:
            p = doc.add_paragraph(style="List Number"); para_spacing(p); add_inline(p, nm.group(2)); i += 1; continue
        if s.startswith(">"):
            p = doc.add_paragraph(); para_spacing(p)
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, s.lstrip("> ").strip(), italic_base=False) if False else add_inline(p, s.lstrip("> ").strip())
            i += 1; continue
        # normal paragraf
        p = doc.add_paragraph(); para_spacing(p); add_inline(p, s)
        i += 1
    if in_code and code_lines:
        add_code_block(code_lines)
    flush_tbl()

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"KAYDEDİLDİ: {OUT}  ({os.path.getsize(OUT)} bytes)")
