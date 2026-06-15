# -*- coding: utf-8 -*-
"""Tüm ders sınav notlarını TEK belgede birleştirir (zero-hallucination: kaynak
docx'ler aynen, yeniden üretilmeden eklenir).

Sıra: BBS656 (sözel) → BBM486 (tasarım örüntüleri) → BYZ652 (yazılım mimarisi).
Her ders kendi öz biçimini korur. Sıra: BBM486 → BBS656.
Çıktı: attached_assets/Tum_Ders_Notlari_Birlesik.docx
"""
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docxcompose.composer import Composer
from docx_helper import new_doc, title_block, h2, _set_run, ACCENT, ACCENT2, GREY

HERE = os.path.dirname(__file__)
ASSETS = os.path.join(HERE, "..", "..", "attached_assets")

NOTES = [
    ("BBM486 — Tasarım Örüntüleri",
     "Sınav Notu (Slide 1–19, Ebru Hoca)",
     os.path.join(ASSETS, "BBM486_Tasarim_Oruntuleri_Sinav_Notu.docx")),
    ("BBS656 — Nesneye Yönelik Yazılım Geliştirme",
     "Sözel Sınav Notu (Slide 0–5, Ebru Hoca)",
     os.path.join(ASSETS, "BBS656_Sozel_Sinav_Notu.docx")),
]

OUT = os.path.join(ASSETS, "BBM486_BBS656_Birlesik_Sinav_Notu.docx")


def make_cover():
    doc = new_doc()
    title_block(
        doc,
        "BBM486 + BBS656 — Birleşik Sınav Notu",
        "BBM486 (Tasarım Örüntüleri) · BBS656 (Nesneye Yönelik Yazılım Geliştirme)",
        "Dr. Nurcan Denli Bayır",
    )
    h2(doc, "İçindekiler")
    for i, (baslik, alt, _) in enumerate(NOTES, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        _set_run(p.add_run(f"{i}.  "), bold=True, color=ACCENT)
        _set_run(p.add_run(baslik), bold=True)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Pt(18)
        _set_run(p2.add_run(alt), italic=True, color=GREY, size=9.5)
    return doc


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def main():
    master = make_cover()
    composer = Composer(master)
    for baslik, alt, path in NOTES:
        if not os.path.exists(path):
            raise FileNotFoundError(path)
        page_break(master)
        composer.append(Document(path))
    composer.save(OUT)
    print("KAYDEDILDI:", os.path.abspath(OUT))


if __name__ == "__main__":
    main()
