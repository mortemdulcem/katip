"""Ortak DOCX biçimlendirme yardımcıları (Times New Roman, temiz başlık hiyerarşisi)."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Times New Roman"
ACCENT = RGBColor(0x1F, 0x3A, 0x5F)      # koyu lacivert
ACCENT2 = RGBColor(0x8B, 0x2E, 0x2E)     # bordo
GREY = RGBColor(0x44, 0x44, 0x44)


def new_doc():
    doc = Document()
    # A4, dar kenar
    for s in doc.sections:
        s.page_width = Cm(21.0)
        s.page_height = Cm(29.7)
        s.left_margin = Cm(1.6)
        s.right_margin = Cm(1.6)
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.12
    pf.space_after = Pt(3)
    return doc


def _set_run(r, size=10.5, bold=False, italic=False, color=None, font=BODY_FONT):
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    rPr = r._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def title_block(doc, title, subtitle=None, meta=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set_run(p.add_run(title), size=20, bold=True, color=ACCENT)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _set_run(p2.add_run(subtitle), size=12.5, bold=True, color=ACCENT2)
    if meta:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(8)
        _set_run(p3.add_run(meta), size=9.5, italic=True, color=GREY)
    _hr(doc)


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F3A5F")
    pbdr.append(bottom)
    pPr.append(pbdr)


def h1(doc, text, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    label = f"{num}  " if num else ""
    _set_run(p.add_run(label + text), size=15, bold=True, color=ACCENT)
    return p


def h2(doc, text, color=ACCENT2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    _set_run(p.add_run(text), size=12, bold=True, color=color)
    return p


def field(doc, label, value):
    """Etiket: değer satırı (kalın etiket, normal değer)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _set_run(p.add_run(f"{label}: "), bold=True, color=ACCENT)
    _set_run(p.add_run(value))
    return p


def para(doc, text, italic=False, color=None, size=10.5):
    p = doc.add_paragraph()
    _set_run(p.add_run(text), italic=italic, color=color, size=size)
    return p


def bullets(doc, items, level=0):
    for it in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(1)
        if isinstance(it, tuple):  # (kalın baş, açıklama)
            _set_run(p.add_run(it[0] + ": "), bold=True)
            _set_run(p.add_run(it[1]))
        else:
            _set_run(p.add_run(it))


def callout(doc, label, text, color=ACCENT2):
    """Vurgu kutusu (kenarlıklı tek paragraf)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "4")
        e.set(qn("w:color"), "8B2E2E")
        pbdr.append(e)
    pPr.append(pbdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "FBF0F0")
    pPr.append(shd)
    _set_run(p.add_run(label + "  "), bold=True, color=color)
    _set_run(p.add_run(text))


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].paragraph_format.space_after = Pt(1)
        _set_run(hdr[i].paragraphs[0].add_run(h), bold=True, size=9.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        _shade_cell(hdr[i], "1F3A5F")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(1)
            _set_run(cells[i].paragraphs[0].add_run(val), size=9.5)
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)
